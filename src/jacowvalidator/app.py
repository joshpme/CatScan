import os

from docx import Document

from flask import Flask, redirect, render_template, request, url_for

from flask_uploads import UploadSet, configure_uploads

from .utils import (
    check_jacow_styles,
    check_margins,
    extract_figures,
    extract_references,
    extract_title,
    get_margins,
    get_page_size,
)

documents = UploadSet("document", ("docx"))

app = Flask(__name__)
app.config.update(
    dict(UPLOADS_DEFAULT_DEST=os.environ.get("UPLOADS_DEFAULT_DEST", "/var/tmp"))
)

configure_uploads(app, (documents,))


@app.template_filter('tick_cross')
def tick_cross(s):
    return "✓" if s else "✗"


@app.route("/")
def hello():
    return redirect(url_for('upload'))


@app.route("/upload", methods=["GET", "POST"])
def upload():
    if request.method == "POST" and documents.name in request.files:
        filename = documents.save(request.files[documents.name])
        fullpath = documents.path(filename)
        try:
            doc = Document(fullpath)

            jacow_styles_ok = check_jacow_styles(doc)

            # get page size and margin details
            sections = []
            for i, section in enumerate(doc.sections):
                sections.append(
                    (
                        get_page_size(section),
                        check_margins(section),
                        get_margins(section),
                    )
                )

            # get title and title syle details
            title = extract_title(doc)

            for i, p in enumerate(doc.paragraphs):
                if p.text.strip().lower() == 'abstract':
                    abstract = {
                        'start': i,
                        'text': p.text,
                        'style': p.style.name,
                        'style_ok': p.style.name in 'JACoW_Abstract_Heading',
                    }
                if p.text.strip().lower() == 'references':
                    references_start = i

            author_paragraphs = doc.paragraphs[1 : abstract['start']]
            authors = {
                'text': ''.join(p.text for p in author_paragraphs),
                'style': set(p.style.name for p in author_paragraphs if p.text.strip()),
                'style_ok': all(
                    p.style.name in ['JACoW_Author List']
                    for p in author_paragraphs
                    if p.text.strip()
                ),
            }

            figures = extract_figures(doc)
            references_in_text, references_list = extract_references(doc)

            return render_template("upload.html", processed=True, **locals())
        finally:
            os.remove(fullpath)

    return render_template("upload.html")
